import sqlite3
from docx import Document
from docx.shared import Pt #（字体大小)
from docx.oxml.ns import qn # 设置字体格式
# import os
# import sys

# sys.path.append(os.path.join(os.path.dirname(os.path.realpath(__file__)), "flush"))
from flush.article import Article

class DataManager:
    '''
    数据库数据操控类
    '''
    def __init__(self, database_path):
        self.database_path = database_path
        self.create_sql = '''create table if not exists {table}(
            number INTEGER primary key autoincrement,
            title VARCHAR not null,
            author VARCHAR not null,
            community VARCHAR not null,
            date VARCHAR not null
            )'''
        self.search_sql = 'select * from {table} where title=?'
        self.insert_sql = 'insert into {table}(title,author,community,date) values(?,?,?,?)'
        self.update_sql = 'update {table} set title=?,author=?,community=?,date=? where title=?'
        self.delete_sql = 'delete from {table} where title=?'
        self.get_table_header_sql = 'select * from {table}'
        self.get_table_data_sql = 'select * from {table}'
        self.get_all_table_name_sql = "SELECT name FROM sqlite_master WHERE type='table'"

    def create_table(self, table):
        '''
        表格若存在则不会覆盖，不存在则新建
        可以在程序`开始运行`时使用
        '''
        con = sqlite3.connect(self.database_path)
        cur = con.cursor()
        try:
            cur.execute(self.create_sql.format(table=table))
            print(f'创建表{table:<30}成功')
        except Exception as e:
            print(e)
            print(f'创建表 {table:<30} 失败')
        finally:
            # 无论如何要关闭表和游标
            cur.close()
            con.close()


    def search_data(self, table, title):
        con = sqlite3.connect(self.database_path)
        cur = con.cursor()

        try:
            cur.execute(self.search_sql.format(table=table), (title, ))
            result = cur.fetchone() #查询不需要提交事务
            return result
        except Exception as e:
            print(e)
            print(f'查询{title:<200}失败')
        finally:
            cur.close()
            con.close()


    def insert_data(self, table, title, author, community, date):
        con = sqlite3.connect(self.database_path)
        cur = con.cursor()

        try:
            cur.execute(self.insert_sql.format(table=table), (title, author, community, date))
            con.commit()
            print(f'插入| {title:<200} |成功')
        except Exception as e:
            print(e)
            con.rollback()
            print(f'插入| {title:<200} |失败')
        finally:
            cur.close()
            con.close()


    def update_data(self, table, title):
        pass


    def delete_data(self, table, title):
        con = sqlite3.connect(self.database_path)
        cur = con.cursor()

        try:
            cur.execute(self.delete_sql.format(table=table), (title, )) # 注意，元组只有一个元素时，必须加逗号
            con.commit()
            print(f'删除{table:<30}中的{title:<200}成功')
        except Exception as e:
            print(e)
            print(f'删除{table:<30}中的{title:<200}失败')
            con.rollback()
        finally:
            cur.close()
            con.close()

    def get_table_header(self, table):
        con = sqlite3.connect(self.database_path)
        cur = con.cursor()

        try:
            cur.execute(self.get_table_header_sql.format(table=table))
            res = cur.fetchall()
            result = [tuple[0] for tuple in res]
            return result
        except Exception as e:
            print(e)
            print(f'获取表头失败！')
        finally:
            cur.close()
            con.close()

    def get_table_data(self, table):
        con = sqlite3.connect(self.database_path)
        cur = con.cursor()

        try:
            cur.execute(self.get_table_data_sql.format(table=table))
            result = cur.fetchall() #查询不需要提交事务
            return result
        except Exception as e:
            print(e)
            print(f'获取表格{table:<200}失败')
        finally:
            cur.close()
            con.close()

    def get_all_table_name(self):
        con = sqlite3.connect(self.database_path)
        cur = con.cursor()

        try:
            cur.execute(self.get_all_table_name_sql)
            result = cur.fetchall()
            result = sorted([i[0] for i in result if i[0] != "sqlite_sequence"], reverse=True)
            return result
        except Exception as e:
            print(e)
            print(f'获取表名失败！')
        finally:
            cur.close()
            con.close()
    
    def export(self, table, filepath, filetype):
        table_data = self.get_table_data(table)
        if filetype == 'Word文档 (*.docx)':
            doc = Document()
            doc.styles["Normal"].font.name=u'Times New Roman'  #设置正文西文字体为Times New Roman
            doc.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')  #设置正文中文字体为宋体
            doc.styles["Normal"].font.size=Pt(12) #设置全局正文为小四（主要是为了列表序号）
            for id, title, author, community, date in table_data:
                article_node = Article(title, author, community, date)

                paragraph1 = doc.add_paragraph(style="List Number")
                paragraph1.paragraph_format.line_spacing = 1.0
                paragraph1.paragraph_format.space_before = 0
                paragraph1.paragraph_format.space_after = 0
                run1 = paragraph1.add_run(f'{article_node.title}')
                run1.font.size = Pt(12)
                run1.bold = True

                paragraph2 = doc.add_paragraph()
                paragraph2.paragraph_format.line_spacing = 1.0
                paragraph2.paragraph_format.space_before = 0
                paragraph2.paragraph_format.space_after = 0
                run2 = paragraph2.add_run(f'{article_node.chinese_title}')
                run2.font.size = Pt(12)
                run2.bold = True

                paragraph3 = doc.add_paragraph()
                paragraph3.paragraph_format.line_spacing = 1.0
                paragraph3.paragraph_format.space_before = 0
                paragraph3.paragraph_format.space_after = 0
                run3 = paragraph3.add_run(f'作者：{article_node.author}')
                run3.font.size = Pt(10.5)
                run3.italic = True

                paragraph4 = doc.add_paragraph()
                paragraph4.paragraph_format.line_spacing = 1.0
                paragraph4.paragraph_format.space_before = 0
                paragraph4.paragraph_format.space_after = 0
                run4 = paragraph4.add_run(f'单位：{article_node.community}')
                run4.font.size = Pt(10.5)
                run4.italic = True

                paragraph5 = doc.add_paragraph()
                paragraph5.paragraph_format.line_spacing = 1.0
                paragraph5.paragraph_format.space_before = 0
                paragraph5.paragraph_format.space_after = 0
                run5 = paragraph5.add_run(f'发表日期：{article_node.date}')
                run5.font.size = Pt(10.5)
                run5.italic = True

                paragraph6 = doc.add_paragraph()
                paragraph6.paragraph_format.line_spacing = 1.0
                paragraph6.paragraph_format.space_before = 0
                paragraph6.paragraph_format.space_after = 0
                run6 = paragraph5.add_run()
                run6.font.size = Pt(10.5)
                run6.italic = True

                # doc.add_paragraph().add_run(f'作者：{article_node.author}\n单位：{article_node.community}\n发表日期：{article_node.date}\n').italic = True
            doc.save(filepath)


        elif filetype == "Text files (*.txt)":
            with open(filepath, 'w', encoding="utf-8") as fp:
                for id, title, author, community, date in table_data:
                    _ = Article(title, author, community, date).format()
                    fp.write(_)
        else:
            raise RuntimeError("filetype 错误！")
        



if __name__ == "__main__":
    database_path = '..\\data\\data.db'
    my_data_manager = DataManager(database_path)
    my_data_manager.create_table('test')
    if my_data_manager.search_data('test', 'hello'):
        print('查询成功')
    else:
        print('查询失败')
    my_data_manager.insert_data('test', 'hello', 'hoven', 'Peking University', '29 June 2022')
    if my_data_manager.search_data('test', 'hello'):
        print('查询成功')
    else:
        print('查询失败')
    print(my_data_manager.get_all_table_name())
    my_data_manager.delete_data('test', 'hi')
    pass